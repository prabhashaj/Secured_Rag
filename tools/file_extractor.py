"""
File Extractor — Extracts text and metadata from PDF, DOCX, TXT, MD, CSV, and JSON files.
"""

from __future__ import annotations

import io
import json
import logging
from typing import Any

import pypdf
import docx

logger = logging.getLogger(__name__)


def extract_text_from_file(file_bytes: bytes, filename: str) -> dict[str, Any]:
    """
    Extract text content, total pages, and metadata from uploaded file bytes.
    Supports PDF (.pdf), Word (.docx), Plain Text (.txt, .md), and Data (.csv, .json).
    """
    ext = filename.lower().split(".")[-1] if "." in filename else ""

    if ext == "pdf":
        return _extract_pdf(file_bytes, filename)
    elif ext == "docx":
        return _extract_docx(file_bytes, filename)
    elif ext in ("txt", "md"):
        return _extract_text(file_bytes, filename)
    elif ext == "json":
        return _extract_json(file_bytes, filename)
    elif ext == "csv":
        return _extract_csv(file_bytes, filename)
    else:
        # Fallback UTF-8 text attempt
        try:
            text = file_bytes.decode("utf-8", errors="ignore")
            return {
                "filename": filename,
                "text": text,
                "pages": 1,
                "file_type": ext or "unknown",
            }
        except Exception as e:
            raise ValueError(f"Unsupported file format '.{ext}': {e}")


def _extract_pdf(file_bytes: bytes, filename: str) -> dict[str, Any]:
    """Extract text page by page from PDF using pypdf."""
    reader = pypdf.PdfReader(io.BytesIO(file_bytes))
    num_pages = len(reader.pages)
    extracted_text = []

    for i, page in enumerate(reader.pages):
        page_text = page.extract_text() or ""
        if page_text.strip():
            extracted_text.append(f"[Page {i + 1}]\n{page_text.strip()}")

    full_text = "\n\n".join(extracted_text)
    if not full_text.strip():
        full_text = f"PDF document {filename} contained no readable text elements."

    return {
        "filename": filename,
        "text": full_text,
        "pages": num_pages,
        "file_type": "pdf",
    }


def _extract_docx(file_bytes: bytes, filename: str) -> dict[str, Any]:
    """Extract text paragraphs and tables from DOCX using python-docx."""
    doc = docx.Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Extract tables if present
    table_texts = []
    for table in doc.tables:
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_cells:
                table_texts.append(" | ".join(row_cells))

    full_text = "\n".join(paragraphs)
    if table_texts:
        full_text += "\n\n=== Document Tables ===\n" + "\n".join(table_texts)

    return {
        "filename": filename,
        "text": full_text or f"DOCX document {filename} contained no text.",
        "pages": max(1, len(paragraphs) // 10),
        "file_type": "docx",
    }


def _extract_text(file_bytes: bytes, filename: str) -> dict[str, Any]:
    """Extract text from TXT or MD files."""
    text = file_bytes.decode("utf-8", errors="replace")
    return {
        "filename": filename,
        "text": text,
        "pages": max(1, len(text.splitlines()) // 30),
        "file_type": filename.split(".")[-1],
    }


def _extract_json(file_bytes: bytes, filename: str) -> dict[str, Any]:
    """Extract formatted text from JSON files."""
    data = json.loads(file_bytes.decode("utf-8"))
    formatted = json.dumps(data, indent=2)
    return {
        "filename": filename,
        "text": formatted,
        "pages": 1,
        "file_type": "json",
    }


def _extract_csv(file_bytes: bytes, filename: str) -> dict[str, Any]:
    """Extract text from CSV files."""
    text = file_bytes.decode("utf-8", errors="replace")
    return {
        "filename": filename,
        "text": text,
        "pages": max(1, len(text.splitlines()) // 40),
        "file_type": "csv",
    }
